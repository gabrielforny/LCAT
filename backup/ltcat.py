import re
import pdfplumber
import win32com.client
from docx import Document
from datetime import datetime
import os
import glob
import threading
from PIL import Image, ImageTk
import time
import requests
import locale
import win32com.client as win32
from docx.shared import Pt
import tkinter as tk
from tkinter import ttk
import pythoncom
import shutil

USERNAME = os.getenv("USERNAME")


# Caminho da pasta
pasta = fr"C:\Users\{USERNAME}\Desktop\ltcat"

# Encontrar todos os arquivos .pdf e .docx
arquivos_pdf = glob.glob(os.path.join(pasta, "*.pdf"))
arquivos_docx = glob.glob(os.path.join(pasta, "*.docx"))

# Verificar se há arquivos para excluir
if not arquivos_pdf and not arquivos_docx:
    print("Nenhum arquivo .pdf ou .docx encontrado na pasta.")
else:
    # Deletar os arquivos encontrados
    for arquivo in arquivos_pdf + arquivos_docx:
        try:
            os.remove(arquivo)
            print(f"Arquivo deletado: {arquivo}")
        except Exception as e:
            print(f"Erro ao excluir {arquivo}: {e}")

    print("Todos os arquivos .pdf e .docx foram excluídos da pasta.")

# Definir o local para o formato brasileiro
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

# Obter a data de hoje
hoje = datetime.now()

ano_atual = datetime.now().year

mes_atual = datetime.now().month

# Formatar a data de hoje no formato "31 de agosto de 2024"
data_hoje = hoje.strftime('%d de %B de %Y')

data_hoje_temp = hoje.strftime('%d-%m-%Y')

# Caminho do arquivo .docx
template_file_path = fr"C:\Users\{USERNAME}\Desktop\arquivos_ltcat\template\template_ltcat_padrao.docx"#r"\\192.168.0.2\tecnico\PGR - GRO\FORMATAÇÃO\TEMPLATE"
pasta_dados = fr"C:\Users\{USERNAME}\Desktop\arquivos_ltcat"#r"\\192.168.0.2\tecnico\PGR - GRO\FORMATAÇÃO\LTCAT"

def mover_arquivos_para_executados():
    try:
        # Diretório de destino para arquivos processados
        pasta_executados = fr"C:\Users\{USERNAME}\Desktop\arquivos_ltcat\executados"

        # Criar a pasta "Executados" se não existir
        if not os.path.exists(pasta_executados):
            os.makedirs(pasta_executados)

        # Mover os arquivos processados
        for arquivo_dados in os.listdir(pasta_dados):
            if arquivo_dados.endswith('.rtf') and not arquivo_dados.startswith('~$'):
                caminho_origem = os.path.join(pasta_dados, arquivo_dados)
                caminho_destino = os.path.join(pasta_executados, arquivo_dados)

                # Mover o arquivo
                shutil.move(caminho_origem, caminho_destino)

                print(f"Arquivo {arquivo_dados} movido para a pasta 'Executados'.")
    except Exception as e:
        print(f"Erro ao mover arquivos: {e}")

def processar_arquivos(progress_label, progress_bar):
    progress_label.config(text="Carregando arquivos da pasta...")
    time.sleep(1)
    
    # Listar todos os arquivos .rtf na pasta, ignorando arquivos temporários (~$)
    arquivos_dados = [f for f in os.listdir(pasta_dados) if f.endswith('.rtf') and not f.startswith('~$')]

    for arquivo_dados in arquivos_dados:
        progress_label.config(text=f"Processando arquivo: {arquivo_dados}...")
        time.sleep(1)

        def format_date(date_str):
            try:
                date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                return date_obj.strftime('%d/%m/%Y')
            except ValueError:
                return 'Data inválida'


        def substituir_variaveis(doc_path, substituicoes):
            doc = Document(doc_path)
            
            progress_label.config(text=f"Processando arquivo: {doc_path}...")
            # Função para criar um novo run e aplicar negrito, fonte e tamanho
            def criar_novo_run(paragrafo, texto, negrito=False, fonte="Verdana", tamanho=8):
                novo_run = paragrafo.add_run(texto)
                novo_run.bold = negrito
                novo_run.font.name = fonte
                novo_run.font.size = Pt(tamanho)
                return novo_run

            # Função para substituir texto em runs de parágrafos ou células de tabela
            def substituir_em_runs(paragrafo, runs, chave, valor):
                progress_label.config(text=f"Preenchendo dados...")
                full_text = ''.join([run.text for run in runs])

                if chave in full_text:
                    # Substitui a chave pelo valor mantendo o resto do texto
                    novo_texto = full_text.replace(chave, valor)

                    # Remove o texto dos runs existentes
                    for run in runs:
                        run.text = ''

                    # Verifica se a chave é "nome_empresa2" para definir fonte e tamanho específicos
                    fonte = "Verdana"
                    tamanho = 16 if chave == "{{nome_empresa2}}" else 8

                    # Recria os runs com o novo texto e aplica negrito ao valor
                    partes = novo_texto.split(valor)
                    if len(partes) == 2:
                        # Primeiro run com o texto antes do valor
                        criar_novo_run(paragrafo, partes[0], fonte=fonte, tamanho=tamanho)

                        # Novo run para o valor com negrito
                        criar_novo_run(paragrafo, valor, negrito=True, fonte=fonte, tamanho=tamanho)

                        # Run com o texto restante
                        criar_novo_run(paragrafo, partes[1], fonte=fonte, tamanho=tamanho)


            # Substituição nos parágrafos
            for p in doc.paragraphs:
                for chave, valor in substituicoes.items():
                    if f"{{{{{chave}}}}}" in p.text:
                        substituir_em_runs(p, p.runs, f"{{{{{chave}}}}}", valor)

            # Substituição nas tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for chave, valor in substituicoes.items():
                                if f"{{{{{chave}}}}}" in p.text:
                                    substituir_em_runs(
                                        p, p.runs, f"{{{{{chave}}}}}", valor)

            # Salva o documento alterado
            doc.save(
                fr'C:\Users\{USERNAME}\Desktop\ltcat\documento_modificado.docx')
            progress_label.config(text="Salvando documento alterado...")

        def format_mes_ano():
            agora = datetime.now()
            mes = agora.strftime('%B').upper()
            ano = agora.strftime('%Y')
            meses_pt = {
                'JANUARY': 'JANEIRO',
                'FEBRUARY': 'FEVEREIRO',
                'MARCH': 'MARÇO',
                'APRIL': 'ABRIL',
                'MAY': 'MAIO',
                'JUNE': 'JUNHO',
                'JULY': 'JULHO',
                'AUGUST': 'AGOSTO',
                'SEPTEMBER': 'SETEMBRO',
                'OCTOBER': 'OUTUBRO',
                'NOVEMBER': 'NOVEMBRO',
                'DECEMBER': 'DEZEMBRO'
            }
            mes_pt = meses_pt.get(mes, mes)
            return f"{mes_pt} DE {ano}"


        # URL e cabeçalhos da API
        url = "https://api.cnpja.com/office/14980183000163"
        headers = {
            "Authorization": "ec1ea1b9-cb4f-460d-8ac1-3cba089fb252-1b1de35e-1616-46e6-9f76-737d7a18194d"
        }

        # Enviar a solicitação GET
        response = requests.get(url, headers=headers)

        if response.status_code == 200:
            data = response.json()
            cnpj = data.get("taxId", "********")
            company = data.get("company", {})
            nome_empresa = company.get("name", "********").upper()
            porte = data['company']['size']['acronym']
            nome_fantasia = data.get("alias", "********").upper()
            data_abertura = data.get("founded", "********").upper()
            data_sit_cad = data.get("statusDate", "********")
            status = data.get("status", {})
            status_text = status.get('text', 'Valor não encontrado').upper()

            nature = company.get('nature', {})
            nature_id = nature.get('id', 'ID não disponível')
            nature_text = nature.get('text', 'Texto não disponível').upper()
            codigo_desc = f"{str(nature_id)[:3]}-{str(nature_id)[-1]} - {nature_text}"

            address = data.get("address", {})
            logradouro = address.get('street', 'Logradouro não disponível').upper()
            numero = address.get('number', 'Número não disponível')
            complemento = address.get('details', 'Complemento não disponível').upper()
            bairro = address.get('district', 'Bairro não disponível').upper()
            municipio = address.get('city', 'Cidade não disponível').upper()
            uf = address.get('state', 'UF não disponível').upper()
            cep = address.get('zip', 'CEP não disponível')

            phones = data.get("phones", [])
            phone_list = [
                f"({telefone.get('area', 'Área não disponível')}) {telefone.get('number', 'Número não disponível')}" for telefone in phones
            ]

            emails = data.get("emails", [])
            email_list = [
                email.get('address', 'Email não disponível').upper() for email in emails]

            main_activity = data.get("mainActivity", {})
            codigo = main_activity.get('id', 'Código não disponível')
            codigo = f"{str(codigo)[:2]}.{str(codigo)[2:4]}-{str(codigo)[4:5]}-{str(codigo)[5:]}"
            atividade = main_activity.get('text', 'Atividade não disponível').upper()
            codigo_completo = f"{codigo} - {atividade}"

            atividade_sec = data.get("sideActivities", [])
            atividade_sec_text = ', '.join(
                atividade_sec) if atividade_sec else "Não informada"

            # Substituições
            replacements = {
                'cnpj': cnpj,
                'nome_empresa': nome_empresa,
                'nome_empresa2':nome_empresa,
                'nomeFantasia': nome_fantasia,
                'dataAbertura': format_date(data_abertura),
                'situacao': status_text,
                'codigoDescricao': codigo_desc,
                'logradouro': logradouro,
                'numero': numero,
                'complemento': complemento,
                'bairro': bairro,
                'municipio': municipio,
                'uf': uf,
                'cep': cep,
                'telefone': ', '.join(phone_list),
                'email': ', '.join(email_list),
                'codigoDescricao': codigo_completo,
                'codigoDescSec': atividade_sec_text,
                'mes_ano': format_mes_ano(),
                'porte': porte,
                'codigo_desc_nat': "*****",
                'dataSitCadastral': format_date(data_sit_cad),
                'situacaoEspecial': "*****",
                'dataSituacaoEsp': "*****",
                'data_hoje': data_hoje,
                'nome_empresarial': nome_empresa,
                'data_hoje_temp': data_hoje_temp
            }
        progress_label.config(text="Preenchendo template...")
        # Exemplo de uso
        substituir_variaveis(
            template_file_path,
            replacements
        )

        # Obter a data de hoje
        hoje = datetime.now()

        data_hoje2 = hoje.strftime('%d/%m/%Y')

        def substituir_marcacoes(doc_path, variaveis, output_path):
            """Substitui variáveis no documento aplicando Verdana 8 para variáveis específicas e mantendo negrito para data_diligencia."""

            doc = Document(doc_path)

            # Lista de variáveis que devem ser formatadas como Verdana 8
            variaveis_verdana8 = {
                "setor2", "atividadeOperacional2", "cargo2", "desc_detal2", "grupo2",
                "agente2", "limiteTolerancia2", "nivelacao2", "meioPropagacao2",
                "frequencia2", "gravidade2", "nivelRisco2", "tempoExposicao2",
                "data2", "medicao2", "tecnicaUtilizada2", "equipamento2",
                "fabricante2", "modelo2", "numSerie2", "dataCalibracao2",
                "fonteGeradora2", "insalubridade2", "periculosidade2",
                "aposentadoria2", "fundamentacaoLegal2", "conclusao2"
            }

            def substituir_texto(paragrafo):
                """Substitui texto dentro de um parágrafo preservando a formatação original,
                aplicando Verdana 8 apenas nas variáveis da lista e negrito para data_diligencia."""

                for chave, valor in variaveis.items():
                    marcador = f'{{{{{chave}}}}}'  # Criando o marcador no formato {{variavel}}

                    if marcador in paragrafo.text:
                        texto_completo = "".join(run.text for run in paragrafo.runs)  # Junta todo o texto do parágrafo

                        if marcador in texto_completo:
                            novo_texto = texto_completo.replace(marcador, valor)  # Substitui a variável

                            # Armazena as propriedades do primeiro run para manter a formatação original
                            fonte_original = paragrafo.runs[0].font.name
                            tamanho_original = paragrafo.runs[0].font.size
                            bold_original = paragrafo.runs[0].bold

                            # Remove os runs existentes para evitar duplicação
                            for run in paragrafo.runs:
                                run.text = ""

                            # Adiciona o texto novamente mantendo a formatação original
                            partes = novo_texto.split(valor)
                            if len(partes) == 2:
                                # Parte antes da variável (mantém a formatação original)
                                if partes[0]:
                                    primeiro_run = paragrafo.add_run(partes[0])
                                    primeiro_run.bold = bold_original
                                    primeiro_run.font.name = fonte_original
                                    primeiro_run.font.size = tamanho_original

                                # Parte da variável (aplica Verdana 8 apenas nas variáveis da lista, e negrito para data_diligencia)
                                novo_run = paragrafo.add_run(valor)
                                if chave == "data_diligencia":
                                    novo_run.bold = True  # Apenas negrito
                                    novo_run.font.name = fonte_original
                                    novo_run.font.size = tamanho_original
                                elif chave in variaveis_verdana8:
                                    novo_run.font.name = "Verdana"
                                    novo_run.font.size = Pt(8)
                                else:
                                    novo_run.font.name = fonte_original
                                    novo_run.font.size = tamanho_original

                                # Parte depois da variável (mantém a formatação original)
                                if partes[1]:
                                    ultimo_run = paragrafo.add_run(partes[1])
                                    ultimo_run.bold = bold_original
                                    ultimo_run.font.name = fonte_original
                                    ultimo_run.font.size = tamanho_original

            # Iterar sobre os parágrafos do documento
            for paragrafo in doc.paragraphs:
                substituir_texto(paragrafo)

            # Iterar sobre as tabelas do documento
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            substituir_texto(paragrafo)

            # Salvar o documento modificado
            doc.save(output_path)
            print(f"Documento salvo com sucesso em {output_path}")


        def rtf_to_pdf(input_file, output_file):
            # Initialize COM for this thread
            pythoncom.CoInitialize()

            try:
                # Inicia o Microsoft Word
                word = win32com.client.gencache.EnsureDispatch('Word.Application')
                word.Visible = False  # Torna invisível para o usuário

                # Abre o arquivo RTF
                doc = word.Documents.Open(os.path.abspath(input_file))

                # Salva como PDF
                output_file_path = os.path.abspath(output_file)
                doc.SaveAs(output_file_path, FileFormat=17)  # 17 é o formato PDF

                doc.Close()  # Fecha o documento
                word.Quit()  # Fecha o aplicativo Word

                print(f"PDF salvo em: {output_file_path}")
            except Exception as e:
                print(f"Erro: {e}")
            finally:
                pythoncom.CoUninitialize()  # Libera o COM



        rtf_to_pdf(
            fr"C:\Users\{USERNAME}\Desktop\arquivos_ltcat\2024 - ABRIL - LTCAT - MAR BRASIL AGENCIA DE VIAGENS E TURISMO LTDA.rtf",
            fr'C:\Users\{USERNAME}\Desktop\ltcat\documento_convertido.pdf'
        )
        progress_label.config(text="Convertendo para PDF...")

        def ler_pdf(caminho_pdf):
            progress_label.config(text="Iniciando leitura do PDF...")
            conclusao2 = ""
            # Abre o PDF
            with pdfplumber.open(caminho_pdf) as pdf:
                
                # Loop através de todas as páginas
                for i, pagina in enumerate(pdf.pages):
                    if conclusao2 != "":
                        break
                    # Extrai o texto de cada página
                    texto = pagina.extract_text()

                    # Expressão regular para capturar a data no formato DD/MM/YYYY
                    padrao = r"Programa de Prevenção de Riscos Ambientais\s+(\d{2}/\d{2}/\d{4})"
                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou a data e exibe o resultado
                    if resultado:
                        vigenciaPpra = resultado.group(1)
                        print(f"Data: {vigenciaPpra}")
                    else:
                        print("Data não encontrada.")

                    # # Expressão regular para capturar o nome da empresa
                    # padrao = r"Empresa\s+(.+)"
                    # resultado = re.search(padrao, texto)

                    # # Verifica se encontrou o nome da empresa e exibe o resultado
                    # if resultado:
                    #     empresa = resultado.group(1).strip()
                    #     print(f"Nome da empresa: {empresa}")
                    # else:
                    #     print("Nome da empresa não encontrado.")

                    # Expressão regular para capturar o endereço e o complemento
                    padrao = r"(Avenida\s[\w\s]+,\s\d+)\s(.+)\s(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2})"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou o endereço e o complemento e exibe o resultado
                    if resultado:
                        endereco2 = resultado.group(1).strip()
                        print(f"Endereço: {endereco2}")
                        complemento2 = resultado.group(2).strip()
                        print(f"Complemento: {complemento2}")
                        cnpj2 = resultado.group(3).strip()
                        print(f"CNPJ: {cnpj2}")
                    else:
                        print("Endereço e complemento não encontrados.")

                    # Expressão regular para capturar CEP, cidade, bairro e UF
                    padrao = r"CEP Cidade Bairro UF\n(\d{5}-\d{3})\s+([\w\s]+?)\s+([\w\s]+?)\s+([A-Z]{2})"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou cada parte e exibe o resultado
                    if resultado:
                        cep2 = resultado.group(1).strip()
                        print(f"CEP: {cep2}")
                        cidade2 = resultado.group(2).strip()
                        print(f"Cidade: {cidade2}")
                        bairro2 = resultado.group(3).strip()
                        print(f"Bairro: {bairro2}")
                        uf2 = resultado.group(4).strip()
                        print(f"UF: {uf2}")

                    else:
                        print("Não foi possível encontrar os dados.")

                    # Expressão regular para capturar CNAE, grau de risco e descrição
                    padrao = r"(\d{4}-\d{1}/\d{2})\s(\d)\s(.+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou o CNAE, grau de risco e descrição e exibe o resultado
                    if resultado:
                        cnae2 = resultado.group(1).strip()
                        print(f"CNAE: {cnae2}")
                        grau_risco2 = resultado.group(2).strip()
                        print(f"Grau de Risco: {grau_risco2}")
                        desc_cnae2 = resultado.group(3).strip()
                        print(f"Descrição CNAE: {desc_cnae2}")

                    else:
                        print("Não foi possível encontrar os dados.")

                    # Expressão regular para capturar os valores de "Previsto" e "Atual"
                    padrao = r"Titular\sSuplente\sDesignado\nPrevisto\s(\d+)\s(\d+)\s(\d+)\nAtual\s(\d+)\s(\d+)\s(\d+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os valores e exibe o resultado
                    if resultado:
                        previsto_titular2 = resultado.group(1).strip()
                        print(f"Titular: {previsto_titular2}")
                        previsto_suplente2 = resultado.group(2).strip()
                        print(f"Titular: {previsto_suplente2}")
                        previsto_designado2 = resultado.group(3).strip()
                        print(f"Designado: {previsto_designado2}")

                        atual_titular2 = resultado.group(4).strip()
                        print(f"  Titular: {atual_titular2}")
                        atual_suplente2 = resultado.group(5).strip()
                        print(f"  Suplente: {atual_suplente2}")
                        atual_designado2 = resultado.group(6).strip()
                        print(f"  Designado: {atual_designado2}")

                    else:
                        print("Não foi possível encontrar os dados.")

                    # Expressão regular para capturar o setor
                    padrao = r"Setor:\s*(.+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou o setor e exibe o resultado
                    if resultado:
                        setor2 = resultado.group(1).strip()
                        print(f"Setor: {setor2}")
                    else:
                        print("Setor não encontrado.")

                    # Expressão regular para capturar "Atividade operacional" antes de "Cargo:"
                    padrao = r"(.+?)\s*Cargo:"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou a atividade e exibe o resultado
                    if resultado:
                        atividadeOperacional2 = resultado.group(1).strip()
                        print(f"Atividade: {atividadeOperacional2}")
                    else:
                        print("Atividade não encontrada.")

                    # Expressão regular para capturar o cargo
                    padrao = r"Cargo:\s*(.+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou o cargo e exibe o resultado
                    if resultado:
                        cargo2 = resultado.group(1).strip()
                        print(f"Cargo: {cargo2}")
                    else:
                        print("Cargo não encontrado.")

                    # Expressão regular para capturar a descrição detalhada
                    padrao = r"Descrição Detalhada:\s*(.*?)\s*Especificação dos Riscos"

                    resultado = re.search(padrao, texto, re.DOTALL)

                    # Verifica se encontrou a descrição detalhada e exibe o resultado
                    if resultado:
                        desc_detal2 = resultado.group(1).strip()
                        print(f"Descrição Detalhada: {desc_detal2}")
                    else:
                        print("Descrição Detalhada não encontrada.")

                    # Expressão regular para capturar o Agente e a Descrição
                    padrao = r"Agente\s+(.+?)\s+\((.+?)\)\s+(.+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os grupos e exibe o resultado
                    if resultado:
                        agente = resultado.group(1).strip()
                        descricao = resultado.group(2).strip()
                        grupo2 = resultado.group(3).strip()
                        print(f"Grupo: {grupo2}")
                        # Combina Agente e Descrição em uma única string
                        agente2 = f"{agente} {descricao}"

                        print(agente2)
                    else:
                        print("Dados não encontrados.")

                    # Expressão regular para capturar Limite de Tolerância e Nível de Ação
                    padrao = r"Limite de Tolerância\s*(\d+,\d+ dB\(A\))\s*Nível de Ação\s*(\d+,\d+ dB\(A\))"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os valores e exibe o resultado
                    if resultado:
                        limite_tolerancia2 = resultado.group(1).strip()
                        print(f"Limite de Tolerância: {limite_tolerancia2}")
                        nivelacao2 = resultado.group(2).strip()
                        print(f"Nível de Ação: {nivelacao2}")
                    else:
                        print("Dados não encontrados.")

                    # Expressão regular para capturar o Meio de Propagação
                    padrao = r"Meio de Propagação\s*(.+?)\."

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou o Meio de Propagação e exibe o resultado
                    if resultado:
                        meioPropagacao2 = resultado.group(1).strip()
                        print(f"Meio de Propagação: {meioPropagacao2}")
                    else:
                        print("Meio de Propagação não encontrado.")

                    # Expressão regular para capturar o texto após "Frequência"
                    padrao = r"Frequência\s*(.+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou o texto e exibe o resultado
                    if resultado:
                        frequencia2 = resultado.group(1).strip()
                        print(f"Exposição Habitual: {frequencia2}")
                    else:
                        print("Texto após 'Frequência' não encontrado.")

                    # Expressão regular para capturar "Gravidade" e o texto associado
                    padrao = r"Gravidade\s*(\w+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou o texto e exibe o resultado
                    if resultado:
                        gravidade2 = resultado.group(1).strip()
                        print(f"Gravidade: {gravidade2}")
                    else:
                        print("Gravidade não encontrada.")

                    # Expressão regular para capturar o texto após "Nível de Risco"
                    padrao = r"Nível de Risco\s*(.+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou o texto e exibe o resultado
                    if resultado:
                        nivelRisco2 = resultado.group(1).strip()
                        print(f"Nível de Risco: {nivelRisco2}")
                    else:
                        print("Nível de Risco não encontrado.")

                    # Expressão regular para capturar o tempo de exposição
                    padrao = r"Tempo de Exposição\s*(\d+h)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou o tempo de exposição e exibe o resultado
                    if resultado:
                        tempoExposicao2 = resultado.group(1).strip()
                        print(f"Tempo de Exposição: {tempoExposicao2}")
                    else:
                        print("Tempo de Exposição não encontrado.")

                    # Expressão regular para capturar a data e o valor de dB(A)
                    padrao = r"(\d{2}/\d{2}/\d{4})\s*(\d+,\d+)\s*dB\(A\)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou a data e o valor de dB(A) e exibe o resultado
                    if resultado:
                        data2 = resultado.group(1).strip()
                        data_diligencia = resultado.group(1).strip()
                        print(f"Data: {data2}")
                        medicao2 = resultado.group(2).strip()
                        print(f"Valor de dB(A): {medicao2}")
                    else:
                        print("Data e valor de dB(A) não encontrados.")

                    # Expressão regular para capturar a data, o valor de dB(A), a empresa, a técnica utilizada e o número de série
                    padrao = r"(\d{2}/\d{2}/\d{4})\s*(\d+,\d+)\s*dB\(A\)\s*([^\:]+):\s*(.+)\s*"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os dados e exibe o resultado
                    if resultado:
                        data2 = resultado.group(1).strip()
                        print(f"Data: {data2}")
                        medicao2 = resultado.group(2).strip()
                        print(f"Valor de dB(A): {medicao2}")
                        tecnicaUtilizada2 = resultado.group(4).strip()
                        print(f"Técnica Utilizada: {tecnicaUtilizada2}")

                    else:
                        print("Dados não encontrados.")

                    # Expressão regular para capturar as partes
                    padrao = r"(\w+)\s+(\w+)\s+(\w+[-\w+]*)\s+(\d+)\s+(\d{2}/\d{2}/\d{4})"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os dados e exibe o resultado
                    if resultado:
                        equipamento = resultado.group(1).strip()
                        print(f"Equipamento: {equipamento}")
                        fabricante = resultado.group(2).strip()
                        print(f"Fabricante: {fabricante}")
                        modelo = resultado.group(3).strip()
                        print(f"Modelo: {modelo}")
                        numSerie2 = resultado.group(4).strip()
                        print(f"Número de Série: {numSerie2}")
                        dataCalibracao2 = resultado.group(5).strip()
                        print(f"Data da Última Calibração: {dataCalibracao2}")
                    else:
                        print("Dados não encontrados.")

                    # Expressão regular para capturar o texto após "Fonte Geradora"
                    padrao = r"Fonte Geradora\s+(.*)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os dados e exibe o resultado
                    if resultado:
                        fonteGeradora2 = resultado.group(1).strip()
                        print(f"Texto após 'Fonte Geradora': {fonteGeradora2}")
                    else:
                        print("Dados não encontrados.")

                    # Expressão regular para capturar as partes
                    padrao = r"Insalubridade\s+(\w+)\s+Periculosidade\s+(\w+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os dados e exibe o resultado
                    if resultado:
                        insalubridade2 = resultado.group(1).strip()
                        print(f"Insalubridade: {insalubridade2}")
                        periculosidade2 = resultado.group(2).strip()
                        print(f"Periculosidade: {periculosidade2}")
                    else:
                        print("Dados não encontrados.")

                    # Expressão regular para capturar as partes
                    padrao = r"Aposentadoria\s+Especial\s+(\w+)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os dados e exibe o resultado
                    if resultado:
                        aposentadoria2 = resultado.group(1).strip()
                        print(f"Aposentadoria Especial: {aposentadoria2}")
                    else:
                        print("Dados não encontrados.")
                    # Expressão regular para capturar as partes
                    padrao = r"Fundamentação Legal\s+([\s\S]*?)\s+Conclusão"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os dados e exibe o resultado
                    if resultado:
                        fundamentacaoLegal2 = resultado.group(1).strip()

                        print(f"Fundamentação Legal:\n{fundamentacaoLegal2}")
                    else:
                        print("Dados não encontrados.")

                    # Expressão regular para capturar a seção "Conclusão"
                    padrao = r"Conclusão\s+([\s\S]*)"

                    resultado = re.search(padrao, texto)

                    # Verifica se encontrou os dados e exibe o resultado
                    if resultado:
                        conclusao2 = resultado.group(1).strip()

                        print(f"Conclusão:\n{conclusao2}")
                    else:
                        print("Dados não encontrados.")

                    variaveis = {
                        'setor2': setor2,
                        'atividadeOperacional2': atividadeOperacional2,
                        'cargo2': cargo2,
                        'desc_detal2': desc_detal2,
                        'grupo2': grupo2,
                        'agente2': agente2,
                        'limiteTolerancia2': limite_tolerancia2,
                        'nivelacao2': nivelacao2,
                        'meioPropagacao2': meioPropagacao2,
                        'frequencia2': frequencia2,
                        'gravidade2': gravidade2,
                        'nivelRisco2': nivelRisco2,
                        'tempoExposicao2': tempoExposicao2,
                        'data2': data2,
                        'data_diligencia': data_diligencia,
                        'medicao2': medicao2,
                        'tecnicaUtilizada2': tecnicaUtilizada2,
                        'equipamento2': equipamento,
                        'fabricante2': fabricante,
                        'modelo2': modelo,
                        'numSerie2': numSerie2,
                        'dataCalibracao2': dataCalibracao2,
                        'fonteGeradora2': fonteGeradora2,
                        'insalubridade2': insalubridade2,
                        'periculosidade2': periculosidade2,
                        'aposentadoria2': aposentadoria2,
                        'fundamentacaoLegal2': fundamentacaoLegal2,
                        'conclusao2': conclusao2
                        # 'data_hoje2': data_hoje2,
                        # 'unidade2': nome_empresa,
                        # #'vigenciaPpra': vigenciaPpra,
                        # 'empresa2': empresa,
                        # 'endereco2': endereco2,
                        # 'complemento2': complemento2,
                        # 'cnpj2': cnpj2,
                        # 'cep2': cep2,
                        # 'cidade2': cidade2,
                        # 'bairro2': bairro2,
                        # 'uf2': uf2,
                        # 'cnae2': cnae2,
                        # 'grau_risco2': grau_risco2,
                        # 'desc_cnae2': desc_cnae2,
                        # 'previsto_titular2': previsto_titular2,
                        # 'previsto_suplente2': previsto_suplente2,
                        # 'previsto_designado2': previsto_designado2,
                        # 'atual_titular2': atual_titular2,
                        # 'atual_suplente2': atual_suplente2,
                        # 'atual_designado2': atual_designado2,
                        
                    }
                    # Substituir as marcações no documento
                    substituir_marcacoes(fr"C:\Users\{USERNAME}\Desktop\ltcat\documento_modificado.docx", variaveis,
                                        f'{mes_atual} {ano_atual} - LTCAT - {nome_empresa}.docx')


        caminho_pdf = fr"C:\Users\{USERNAME}\Desktop\ltcat\documento_convertido.pdf"
        ler_pdf(caminho_pdf)
        progress_label.config(text="Salvando arquivo como PDF...")


        def atualizar_indice(doc_path):
            # Initialize COM for this thread
            pythoncom.CoInitialize()
            # Inicializar o Word
            word = win32.Dispatch("Word.Application")
            word.Visible = False  # Define como False para não abrir o Word visivelmente

            # Abrir o documento
            doc = word.Documents.Open(doc_path)

            # Atualizar os índices/tabelas de conteúdo (sumário)
            for toc in doc.TablesOfContents:
                toc.Update()  # Atualiza o índice sem atualizar os números de página
                toc.UpdatePageNumbers()  # Atualiza os números de página

            # Salvar e fechar o documento
            doc.Save()
            doc.Close()

            # Fechar o Word
            word.Quit()

        # Atualizar índice
        atualizar_indice(
            fr"C:\Users\{USERNAME}\Desktop\ltcat\documento_modificado.docx")
        
        time.sleep(2)
        
        rtf_to_pdf(
            fr'C:\Users\{USERNAME}\Desktop\ltcat\documento_modificado.docx',
            fr'C:\Users\{USERNAME}\Desktop\ltcat\{mes_atual} {ano_atual} - LTCAT - {nome_empresa}.pdf')
        # Mover os arquivos processados para a pasta "Executados"
        mover_arquivos_para_executados()

    progress_label.config(text="Processo concluído!")
    progress_bar.stop()
#Função para iniciar a execução em uma thread separada
def start_process():
    progress_bar.start()
    threading.Thread(target=processar_arquivos, args=(
        progress_label, progress_bar)).start()

# Interface gráfica com Tkinter
root = tk.Tk()
root.title("Processar Arquivos LTCAT")
root.geometry("400x300")

# Logo da empresa
logo_image = Image.open(fr"C:\Users\{USERNAME}\Desktop\ltcat\logo_empresa.jpg")
logo_image = logo_image.resize((200, 100), Image.LANCZOS)
logo_photo = ImageTk.PhotoImage(logo_image)
logo_label = tk.Label(root, image=logo_photo)
logo_label.pack(pady=10)

# Botão para processar arquivos
botao_processar = tk.Button(
    root, text="Processar arquivos LTCAT", command=start_process)
botao_processar.pack(pady=10)

# Barra de progresso
progress_bar = ttk.Progressbar(
    root, orient="horizontal", mode="indeterminate", length=280)
progress_bar.pack(pady=10)

# Label de status do processo
progress_label = tk.Label(root, text="Aguardando...")
progress_label.pack()

# Iniciar a interface Tkinter
root.mainloop()