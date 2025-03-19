from docx import Document
from docx.enum.text import WD_BREAK

def ajustar_conclusao_no_docx(caminho_arquivo):
    # Carrega o documento
    doc = Document(caminho_arquivo)
    
    # Define o texto que marca o início da seção 9
    texto_busca = "CONCLUSÃO DO LAUDO TÉCNICO DAS CONDIÇÕES AMBIENTAIS DO TRABALHO – LTCAT"
    
    for i, paragrafo in enumerate(doc.paragraphs):
        if texto_busca in paragrafo.text:
            novo_paragrafo = doc.paragraphs[i]._element
            run = paragrafo.insert_paragraph_before().add_run()
            run.add_break(WD_BREAK.PAGE)
            break  # Sai do loop após encontrar e modificar
    
    # Salva o novo documento
    novo_caminho = caminho_arquivo.replace(".docx", "_ajustado.docx")
    doc.save(novo_caminho)
    print(f"Documento ajustado salvo em: {novo_caminho}")

# Exemplo de uso
ajustar_conclusao_no_docx(f"C:\\Users\\Gabriel\\tecnico\\PGR-GRO\\00 - RENOVADOS 2024\\EXECUTADOS LTCAT\\3 2025 - LTCAT - MAR BRASIL AGENCIA DE VIAGENS E TURISMO LTDA.docx")
